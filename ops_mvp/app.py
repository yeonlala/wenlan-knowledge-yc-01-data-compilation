# -*- coding: utf-8 -*-
"""
yc-checking 运维可视化：脚本目录与流水线一键运行（三阶段主线）。

不修改仓库根目录下的业务脚本，仅通过 subprocess 调用。

启动（在仓库根 yc-checking 下）：
  uvicorn ops_mvp.app:app --host 127.0.0.1 --port 8765 --reload
"""

from __future__ import annotations

import json
import os
import shutil
import subprocess
import sys
import tempfile
import time
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from urllib.parse import quote

from docx import Document
from fastapi import FastAPI, File, Form, HTTPException, Request, UploadFile
from starlette.background import BackgroundTask
from fastapi.responses import FileResponse, HTMLResponse, PlainTextResponse
from fastapi.templating import Jinja2Templates
from pydantic import BaseModel, Field

from ops_mvp.discovery import (
    build_filename_order_hint,
    filter_core_chain_lines,
    filter_data_flow_steps,
    filter_pipeline_catalog,
    playbook_step_runnable,
)
from ops_mvp.pipeline_catalog import PipelineScript, get_script_by_id
from ops_mvp.project_guide import (
    CORE_CHAIN_LINE_ITEMS,
    DATA_FLOW_STEPS,
    OPERATION_PRINCIPLE,
    PROJECT_SUMMARY,
    PROJECT_TITLE,
    WHO_USES,
)
from ops_mvp.workspace import build_workspace_snapshot

REPO_ROOT = Path(__file__).resolve().parent.parent
MARKDOWN_SCRIPTS = {
    "check": REPO_ROOT / "4_markdown_quality_checker.py",
    "fix": REPO_ROOT / "5_markdown_quality_fix.py",
}

templates = Jinja2Templates(directory=str(Path(__file__).resolve().parent / "templates"))

MAX_JSON_READ_BYTES = 2 * 1024 * 1024
MAX_PREVIEW_BINARY_BYTES = 20 * 1024 * 1024
MAX_DOWNLOAD_BYTES = 80 * 1024 * 1024
MAX_FOLDER_ZIP_BYTES = 512 * 1024 * 1024
MAX_FIRST_BATCH_IMPORT_BYTES = 512 * 1024 * 1024
MAX_TREE_NODES = 350

app = FastAPI(title="yc-checking 运维台", docs_url=None, redoc_url=None)

# 推荐顺序：三阶段 ①→②→③（缺脚本时由 discovery 过滤）
OPERATIONS_PLAYBOOK: List[Dict[str, Any]] = [
    {
        "step": 1,
        "badge": "①",
        "title": "阶段一 · 资料包验收（可选）",
        "detail": "扫描「第一批」→ 检查结果。",
        "script_id": "check_tobacco_kb",
    },
    {
        "step": 2,
        "badge": "②",
        "title": "阶段二 · 提取三份清单",
        "detail": "确认单 → extract_jsons 下三份清单 JSON + docs_metadata。",
        "script_id": "export_confirmation_bundle",
    },
    {
        "step": 3,
        "badge": "③",
        "title": "阶段三 · 提取文字",
        "detail": "正文抽取 → kb_local/extracted_markdown + manifest。",
        "script_id": "prepare_local_kb",
    },
]

# 流水线分组在界面中的展示顺序（未出现在此列表的组名排在后面）
CATALOG_GROUP_ORDER: List[str] = [
    "资料包验收",
    "清单与归集",
    "正文抽取",
    "质量与安全",
    "工具与测试",
    "维护",
]


def _markdown_paths() -> Dict[str, Path]:
    return dict(MARKDOWN_SCRIPTS)


def _root_resolved() -> Path:
    return REPO_ROOT.resolve()


def safe_relative_path(user: str) -> Path:
    if not user or not user.strip():
        return _root_resolved()
    raw = user.strip().replace("\\", "/")
    if ".." in Path(raw).parts:
        raise HTTPException(status_code=400, detail="路径不允许包含 ..")
    p = (_root_resolved() / raw).resolve()
    try:
        p.relative_to(_root_resolved())
    except ValueError as e:
        raise HTTPException(status_code=400, detail="路径必须位于项目仓库内") from e
    return p


def _extract_docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts: List[str] = []
    for para in doc.paragraphs:
        parts.append(para.text or "")
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").replace("\n", " ").strip() for c in row.cells]
            parts.append("\t".join(cells))
    out = "\n".join(parts).strip()
    return out if out else "（空文档或无可提取段落）"


def _extract_legacy_doc_text(path: Path) -> Optional[str]:
    antiword = shutil.which("antiword")
    if antiword:
        try:
            r = subprocess.run(
                [antiword, str(path.resolve())],
                capture_output=True,
                timeout=120,
            )
            if r.returncode == 0 and r.stdout:
                return r.stdout.decode("utf-8", errors="replace")
        except (OSError, subprocess.SubprocessError):
            pass

    for cmd in ("soffice", "libreoffice"):
        exe = shutil.which(cmd)
        if not exe:
            continue
        try:
            with tempfile.TemporaryDirectory() as td:
                td_path = Path(td)
                r = subprocess.run(
                    [
                        exe,
                        "--headless",
                        "--convert-to",
                        "txt:Text",
                        "--outdir",
                        str(td_path),
                        str(path.resolve()),
                    ],
                    capture_output=True,
                    timeout=180,
                )
                if r.returncode != 0:
                    continue
                direct = td_path / f"{path.stem}.txt"
                if direct.is_file():
                    return direct.read_text(encoding="utf-8", errors="replace")
                txts = list(td_path.glob("*.txt"))
                if len(txts) == 1:
                    return txts[0].read_text(encoding="utf-8", errors="replace")
        except (OSError, subprocess.SubprocessError):
            continue
    return None


class RunBody(BaseModel):
    scan_root: str = Field(default="extract_jsons", description="待扫描的相对路径")
    checker_out: str = Field(default="markdown_quality_reports", description="检测报告输出目录")
    fix_out: str = Field(default="extract_jsons_fixed", description="修复副本输出目录")


class RunResult(BaseModel):
    ok: bool
    action: str
    exit_code: int
    stdout: str
    stderr: str
    duration_ms: int
    command: List[str]


class ChainStepRecord(BaseModel):
    script_id: str
    skipped: bool = False
    skip_reason: Optional[str] = None
    ok: bool = True
    exit_code: int = 0
    stdout: str = ""
    stderr: str = ""
    duration_ms: int = 0
    command: List[str] = Field(default_factory=list)


class ChainRunResponse(BaseModel):
    ok: bool
    stopped_at: Optional[str] = None
    steps: List[ChainStepRecord]


class CleanRunBody(BaseModel):
    dry_run: bool = Field(default=True, description="True 为仅预览；False 为真实删除（须前端二次确认）")


class FirstBatchImportPathBody(BaseModel):
    source_path: str = Field(..., description="本机目录的绝对路径")
    target_name: Optional[str] = Field(
        default=None,
        description="写入仓库「第一批」下的文件夹名；默认使用源目录名",
    )


MOCK_SCENARIO_CHOICES = frozenset(
    {
        "full",
        "dirs_only",
        "missing_required_partial",
    }
)


class MockGenerateBody(BaseModel):
    scenario: str = Field(
        default="full",
        description="与 generate_mock_tobacco_project.py --scenario 一致：full / dirs_only / missing_required_partial",
    )
    project_name: Optional[str] = Field(
        default=None,
        description="mockdata 下项目文件夹名；省略则脚本按客户/项目段自动生成",
    )
    customer: Optional[str] = Field(
        default=None,
        description="与包内七段文件名第一段的客户段一致；省略用脚本默认「测试市局」",
    )
    project: Optional[str] = Field(
        default=None,
        description="文件名第二段项目简称（勿含下划线）；省略用脚本默认「知识库验收Mock」",
    )
    seed: Optional[int] = Field(
        default=None,
        description="missing_required_partial 时可选随机种子，便于复现",
    )


class WorkspaceDeleteBody(BaseModel):
    relative: str = Field(..., description="相对仓库根的路径（文件或文件夹）")


def _assert_workspace_delete_allowed(target: Path, root: Path) -> None:
    rt = target.resolve()
    rr = root.resolve()
    if rt == rr:
        raise HTTPException(status_code=400, detail="禁止删除仓库根目录")
    try:
        rt.relative_to(rr)
    except ValueError as e:
        raise HTTPException(status_code=400, detail="路径不在仓库内") from e
    gd = (root / ".git").resolve()
    if gd.exists():
        try:
            rt.relative_to(gd)
            raise HTTPException(status_code=400, detail="禁止删除版本库目录 .git")
        except ValueError:
            pass


def _assert_extract_jsons_zip_folder(target: Path, root: Path) -> None:
    """仅允许将 extract_jsons 下的子目录（含嵌套）打包为 zip，禁止整包根目录 extract_jsons。"""
    _assert_workspace_delete_allowed(target, root)
    if not target.is_dir():
        raise HTTPException(status_code=400, detail="只能打包目录")
    rel = str(target.relative_to(root)).replace("\\", "/")
    if rel == "extract_jsons":
        raise HTTPException(
            status_code=400,
            detail="不能打包整个 extract_jsons 根目录，请进入后选择具体项目文件夹",
        )
    parts = rel.split("/")
    if len(parts) < 2 or parts[0] != "extract_jsons":
        raise HTTPException(
            status_code=400,
            detail="仅支持打包 extract_jsons 下的子文件夹",
        )


def _allocate_dest_under(parent: Path, folder_name: str) -> Tuple[Path, str]:
    """在 parent 下分配一个不存在的子目录路径；若重名则自动加 _2、_3…"""
    dest = parent / folder_name
    if not dest.exists():
        return dest, folder_name
    for i in range(2, 10000):
        alt = f"{folder_name}_{i}"
        cand = parent / alt
        if not cand.exists():
            return cand, alt
    raise HTTPException(status_code=500, detail="无法分配可用文件夹名")


def _parts_from_browser_upload_filename(fn: str) -> Tuple[str, ...]:
    """浏览器文件夹上传时 multipart 内嵌的相对路径（webkitRelativePath）。"""
    if not fn:
        raise HTTPException(status_code=400, detail="缺少文件相对路径")
    n = fn.replace("\\", "/").strip().lstrip("/")
    parts = tuple(Path(n).parts)
    if not parts:
        raise HTTPException(status_code=400, detail="路径无效")
    if ".." in parts:
        raise HTTPException(status_code=400, detail="路径非法")
    return parts


def _safe_first_batch_folder_name(name: str) -> str:
    s = (name or "").strip()
    if not s:
        raise HTTPException(status_code=400, detail="文件夹名为空")
    if len(s) > 180:
        raise HTTPException(status_code=400, detail="文件夹名过长")
    if "\x00" in s:
        raise HTTPException(status_code=400, detail="文件夹名含有非法字符")
    if any(x in s for x in ("/", "\\", ":", "<", ">", "|", "?", "*")):
        raise HTTPException(
            status_code=400,
            detail="文件夹名不能含路径分隔符或非法符号",
        )
    if s in (".", ".."):
        raise HTTPException(status_code=400, detail="非法文件夹名")
    return s


def _run_subprocess(script_path: Path, args: List[str]) -> RunResult:
    if not script_path.is_file():
        raise HTTPException(
            status_code=503,
            detail=f"未找到脚本文件：{script_path.name}（请放在仓库根目录，与运维台说明一致）",
        )
    cmd = [sys.executable, str(script_path.name)] + args
    t0 = time.perf_counter()
    proc = subprocess.run(
        cmd,
        cwd=str(REPO_ROOT),
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        env={**os.environ, "PYTHONUTF8": "1"},
        timeout=7200,
    )
    ms = int((time.perf_counter() - t0) * 1000)
    name = script_path.name
    return RunResult(
        ok=proc.returncode == 0,
        action=name,
        exit_code=proc.returncode,
        stdout=proc.stdout or "",
        stderr=proc.stderr or "",
        duration_ms=ms,
        command=cmd,
    )


def _script_to_dict(s: PipelineScript) -> Dict[str, Any]:
    return {
        "id": s.id,
        "filename": s.filename,
        "title": s.title,
        "group": s.group,
        "description": s.description,
        "runner": s.runner,
        "phase": s.phase,
        "when_use": s.when_use,
        "inputs_desc": s.inputs_desc,
        "outputs_desc": s.outputs_desc,
        "default_argv": list(s.default_argv),
        "cli_example": s.cli_example,
        "warn": s.warn,
    }


def _order_catalog_groups(by_group: Dict[str, List[Dict[str, Any]]]) -> Dict[str, List[Dict[str, Any]]]:
    ordered: Dict[str, List[Dict[str, Any]]] = {}
    for name in CATALOG_GROUP_ORDER:
        if name in by_group:
            ordered[name] = by_group[name]
    for name, rows in by_group.items():
        if name not in ordered:
            ordered[name] = rows
    return ordered


@app.get("/", response_class=HTMLResponse)
def index(request: Request) -> Any:
    return templates.TemplateResponse(
        "index.html",
        {
            "request": request,
            "repo_root": str(REPO_ROOT),
            "default_scan": "extract_jsons",
            "default_checker_out": "markdown_quality_reports",
            "default_fix_out": "extract_jsons_fixed",
        },
    )


@app.get("/api/workspace/overview")
def api_workspace_overview() -> Dict[str, Any]:
    """关键目录、产物计数、extract_jsons 项目、检查结果近期文件、根目录脚本列表。"""
    snap = build_workspace_snapshot(REPO_ROOT)
    snap["repo_root"] = str(REPO_ROOT)
    md = _markdown_paths()
    kd = snap.get("key_directories") or []
    if not md["check"].is_file():
        kd = [x for x in kd if x.get("relative") != "markdown_quality_reports"]
    if not md["fix"].is_file():
        kd = [x for x in kd if x.get("relative") != "extract_jsons_fixed"]
    snap["key_directories"] = kd
    return snap


@app.get("/api/help/playbook")
def api_help_playbook() -> Dict[str, Any]:
    md = _markdown_paths()
    steps_out: List[Dict[str, Any]] = []
    for row in OPERATIONS_PLAYBOOK:
        if not playbook_step_runnable(REPO_ROOT, row["script_id"], md):
            continue
        steps_out.append({**row, "step": len(steps_out) + 1})
    return {"title": "推荐执行顺序", "steps": steps_out}


def _guide_payload() -> Dict[str, Any]:
    """项目说明 JSON：按仓库内实际存在的脚本裁剪条目。"""
    md = _markdown_paths()
    return {
        "title": PROJECT_TITLE,
        "summary": PROJECT_SUMMARY,
        "who_uses": WHO_USES,
        "operation_principle": OPERATION_PRINCIPLE,
        "core_chain_line_items": filter_core_chain_lines(CORE_CHAIN_LINE_ITEMS, md),
        "data_flow_steps": filter_data_flow_steps(REPO_ROOT, md, DATA_FLOW_STEPS),
        "filename_order_hint": build_filename_order_hint(REPO_ROOT, md),
    }


@app.get("/api/guide")
@app.get("/api/guide/")
@app.get("/api/help/guide")
def api_guide() -> Dict[str, Any]:
    """项目目的、数据流说明。别名：`/api/guide/`、`/api/help/guide`。"""
    return _guide_payload()


@app.get("/api/catalog")
def api_catalog() -> Dict[str, Any]:
    md = _markdown_paths()
    available = filter_pipeline_catalog(REPO_ROOT, md)
    by_group: Dict[str, List[Dict[str, Any]]] = {}
    for s in available:
        by_group.setdefault(s.group, []).append(_script_to_dict(s))
    ordered = _order_catalog_groups(by_group)
    return {
        "repo_root": str(REPO_ROOT),
        "groups": ordered,
        "scripts": [_script_to_dict(s) for s in available],
        "group_order": CATALOG_GROUP_ORDER,
    }


@app.get("/api/info")
def api_info() -> Dict[str, Any]:
    md = _markdown_paths()
    ck = md["check"].is_file()
    fx = md["fix"].is_file()
    return {
        "repo_root": str(REPO_ROOT),
        "markdown_scripts": {k: str(v) for k, v in MARKDOWN_SCRIPTS.items()},
        "guide_urls": ["/api/guide", "/api/guide/", "/api/help/guide"],
        "features": {
            "markdown_check": ck,
            "markdown_fix": fx,
            "markdown_pipeline": ck or fx,
        },
    }


@app.post("/api/run/check")
def api_run_check(body: RunBody) -> RunResult:
    scan = safe_relative_path(body.scan_root)
    if not scan.exists():
        raise HTTPException(status_code=400, detail=f"路径不存在：{scan}")
    rel_scan = scan.relative_to(_root_resolved())
    out = safe_relative_path(body.checker_out)
    rel_out = out.relative_to(_root_resolved())
    args = [str(rel_scan), "--out", str(rel_out)]
    return _run_subprocess(MARKDOWN_SCRIPTS["check"], args)


@app.post("/api/run/fix")
def api_run_fix(body: RunBody, dry_run: bool = False) -> RunResult:
    scan = safe_relative_path(body.scan_root)
    if not scan.exists():
        raise HTTPException(status_code=400, detail=f"路径不存在：{scan}")
    rel_scan = scan.relative_to(_root_resolved())
    out = safe_relative_path(body.fix_out)
    rel_out = out.relative_to(_root_resolved())
    args = [str(rel_scan), "--out-dir", str(rel_out)]
    if dry_run:
        args.append("--dry-run")
    return _run_subprocess(MARKDOWN_SCRIPTS["fix"], args)


def _invoke_pipeline_script(script_id: str, body: RunBody) -> RunResult:
    """执行单个流水线脚本（与 /api/run/pipeline/{id} 行为一致）。"""
    if script_id == "markdown_check":
        return api_run_check(body)
    if script_id == "markdown_fix":
        return api_run_fix(body, dry_run=False)

    meta = get_script_by_id(script_id)
    if not meta:
        raise HTTPException(status_code=404, detail="未知脚本 id")
    if meta.runner != "subprocess_argv":
        raise HTTPException(
            status_code=400,
            detail="该脚本不在面板内一键运行范围，请使用命令行或复制示例命令。",
        )
    script_path = REPO_ROOT / meta.filename
    argv = [str(x) for x in meta.default_argv]
    return _run_subprocess(script_path, argv)


@app.post("/api/run/pipeline/{script_id}")
def api_run_pipeline(script_id: str, body: RunBody) -> RunResult:
    """一键流水线；Markdown 类 id 与脚本清单中的 id 对齐，共用表单路径参数。"""
    return _invoke_pipeline_script(script_id, body)


@app.post("/api/run/chain/main", response_model=ChainRunResponse)
def api_run_chain_main(body: RunBody) -> ChainRunResponse:
    """
    一键串跑：**①→②→③**（验收 → 三份清单 → 提取文字）；不含 Markdown 质检/修复。
    仓库内不存在的脚本会跳过；任一步失败则中止并返回已执行步骤的输出。
    """
    scan = safe_relative_path(body.scan_root)
    if not scan.exists():
        raise HTTPException(status_code=400, detail=f"路径不存在：{scan}")

    md = _markdown_paths()
    available = {s.id for s in filter_pipeline_catalog(REPO_ROOT, md)}
    order: List[str] = [
        "check_tobacco_kb",
        "export_confirmation_bundle",
        "prepare_local_kb",
    ]

    steps: List[ChainStepRecord] = []
    for sid in order:
        if sid not in available:
            steps.append(
                ChainStepRecord(
                    script_id=sid,
                    skipped=True,
                    skip_reason="仓库内无此脚本",
                    ok=True,
                )
            )
            continue
        try:
            r = _invoke_pipeline_script(sid, body)
        except HTTPException as e:
            detail = e.detail
            if not isinstance(detail, str):
                detail = str(detail)
            steps.append(
                ChainStepRecord(
                    script_id=sid,
                    skipped=False,
                    ok=False,
                    exit_code=1,
                    stdout="",
                    stderr=detail,
                    duration_ms=0,
                    command=[],
                )
            )
            return ChainRunResponse(ok=False, stopped_at=sid, steps=steps)

        steps.append(
            ChainStepRecord(
                script_id=sid,
                skipped=False,
                ok=r.ok,
                exit_code=r.exit_code,
                stdout=r.stdout,
                stderr=r.stderr,
                duration_ms=r.duration_ms,
                command=r.command,
            )
        )
        if not r.ok:
            return ChainRunResponse(ok=False, stopped_at=sid, steps=steps)

    return ChainRunResponse(ok=True, stopped_at=None, steps=steps)


@app.post("/api/run/clean")
def api_run_clean(body: CleanRunBody) -> RunResult:
    """
    运维台一键清理：清空「检查结果」与「extract_jsons」内全部条目（目录本身保留）。
    等价：`clean_workspace.py --results --extract-jsons` + `--dry-run` 或不带预览。
    """
    script_path = REPO_ROOT / "clean_workspace.py"
    argv = ["--results", "--extract-jsons"]
    if body.dry_run:
        argv.append("--dry-run")
    return _run_subprocess(script_path, argv)


@app.post("/api/first-batch/import-path")
def api_first_batch_import_path(body: FirstBatchImportPathBody) -> Dict[str, Any]:
    """从本机目录复制单个项目到仓库「第一批/<文件夹名>」（运维台与仓库同机时使用）。"""
    src = Path(body.source_path).expanduser()
    try:
        src = src.resolve()
    except OSError as e:
        raise HTTPException(status_code=400, detail=f"无法解析路径：{e}") from e
    if not src.is_dir():
        raise HTTPException(status_code=400, detail="源路径不是有效目录")

    root = _root_resolved()
    name_src = _safe_first_batch_folder_name(src.name)
    final_name = (
        _safe_first_batch_folder_name(body.target_name)
        if body.target_name and body.target_name.strip()
        else name_src
    )
    planned = root / "第一批" / final_name
    rs = src.resolve()
    rp = planned.resolve()
    if rs == rp:
        raise HTTPException(status_code=400, detail="源目录与目标相同")
    try:
        rp.relative_to(rs)
    except ValueError:
        pass
    else:
        raise HTTPException(
            status_code=400,
            detail="目标将位于源目录内部，已拒绝（请更换目标文件夹名或源路径）",
        )

    (root / "第一批").mkdir(parents=True, exist_ok=True)
    dest, used_name = _allocate_dest_under(root / "第一批", final_name)
    shutil.copytree(src, dest)
    hint = "已从本机路径复制到「第一批」"
    if used_name != final_name:
        hint += "（同名已存在，已自动加后缀）"
    return {"ok": True, "target": f"第一批/{used_name}", "hint": hint}


@app.post("/api/first-batch/import-directory")
async def api_first_batch_import_directory(
    files: List[UploadFile] = File(...),
    project_name: Optional[str] = Form(None),
) -> Dict[str, Any]:
    """
    浏览器「选择文件夹」（webkitdirectory）上传：去掉所选目录名的最外层后，
    将内部相对路径写入「第一批/…」（所选文件夹本身不作为「第一批」下的一层目录）。
    若表单传入「目标文件夹名」，则先落在「第一批/<该名>/」之下（多一层归集夹）。
    空子文件夹无法随浏览器上传，请使用下方「本机路径」整目录复制。
    """
    if not files:
        raise HTTPException(
            status_code=400,
            detail="请选择文件夹（至少包含一个文件；纯空文件夹浏览器不会上传）",
        )

    roots: set[str] = set()
    planned: List[Tuple[UploadFile, Tuple[str, ...]]] = []
    for uf in files:
        parts = _parts_from_browser_upload_filename(uf.filename or "")
        roots.add(parts[0])
        planned.append((uf, parts))

    if len(roots) != 1:
        raise HTTPException(
            status_code=400,
            detail="一次只能选一个顶层项目文件夹（请勿多选并列目录；可只选该项目文件夹本身）",
        )

    repo_root = _root_resolved()
    first_batch = repo_root / "第一批"
    first_batch.mkdir(parents=True, exist_ok=True)
    fb_resolved = first_batch.resolve()

    user_named = bool(project_name and project_name.strip())
    requested_top = ""
    if user_named:
        requested_top = _safe_first_batch_folder_name(project_name.strip())
        dest_base, used_top = _allocate_dest_under(first_batch, requested_top)
    else:
        dest_base = first_batch
        used_top = ""

    total_bytes = 0
    file_count = 0
    written_paths: List[Path] = []
    try:
        for uf, parts in planned:
            tail = parts[1:]
            if not tail:
                raise HTTPException(
                    status_code=400,
                    detail="相对路径缺少文件部分，请检查所选文件夹或使用「本机路径」导入",
                )
            rel_tail = Path(*tail)
            out_path = (dest_base / rel_tail).resolve()
            try:
                out_path.relative_to(dest_base.resolve())
            except ValueError as e:
                raise HTTPException(status_code=400, detail="路径非法") from e
            try:
                out_path.relative_to(fb_resolved)
            except ValueError as e:
                raise HTTPException(status_code=400, detail="路径非法") from e

            raw = await uf.read()
            total_bytes += len(raw)
            if total_bytes > MAX_FIRST_BATCH_IMPORT_BYTES:
                raise HTTPException(status_code=413, detail="上传总体积超过上限（约 512MB）")
            out_path.parent.mkdir(parents=True, exist_ok=True)
            out_path.write_bytes(raw)
            written_paths.append(out_path)
            file_count += 1
    except HTTPException:
        for p in reversed(written_paths):
            try:
                if p.is_file():
                    p.unlink()
            except OSError:
                pass
        raise
    except Exception as e:
        for p in reversed(written_paths):
            try:
                if p.is_file():
                    p.unlink()
            except OSError:
                pass
        raise HTTPException(status_code=500, detail=f"导入失败：{e}") from e

    hint = f"已从浏览器上传 {file_count} 个文件到「第一批」"
    if user_named and used_top != requested_top:
        hint += "（「目标文件夹名」已自动加后缀避免重名）"
    hint += "（所选顶层文件夹名不会作为单独一层目录写入）。"
    hint += "空子文件夹不会被浏览器上传，需要完整目录时请用「本机路径」复制。"
    target_str = f"第一批/{used_top}" if user_named else "第一批"
    return {
        "ok": True,
        "target": target_str,
        "files_written": file_count,
        "hint": hint,
    }


@app.post("/api/run/generate-mock")
def api_run_generate_mock(body: MockGenerateBody) -> RunResult:
    """在仓库「mockdata」下运行 generate_mock_tobacco_project.py（运维台一键生成模拟资料）。"""
    scenario = (body.scenario or "full").strip()
    if scenario not in MOCK_SCENARIO_CHOICES:
        raise HTTPException(status_code=400, detail=f"无效 scenario：{scenario}")
    script_path = REPO_ROOT / "generate_mock_tobacco_project.py"
    argv: List[str] = ["--batch", "mockdata", "--scenario", scenario]
    pn = (body.project_name or "").strip()
    if pn:
        argv.extend(["--project-name", pn])
    cust = (body.customer or "").strip()
    if cust:
        argv.extend(["--customer", cust])
    proj = (body.project or "").strip()
    if proj:
        argv.extend(["--project", proj])
    if body.seed is not None:
        argv.extend(["--seed", str(int(body.seed))])
    return _run_subprocess(script_path, argv)


@app.get("/api/reports/dashboard")
def api_reports_dashboard(
    checker_out: str = "markdown_quality_reports",
    fix_out: str = "extract_jsons_fixed",
) -> Dict[str, Any]:
    root = _root_resolved()
    co = safe_relative_path(checker_out)
    fo = safe_relative_path(fix_out)

    summary_path = co / "summary.json"
    fix_summary_path = fo / "fix_summary.json"

    checker_summary: Any = None
    if summary_path.is_file():
        if summary_path.stat().st_size > MAX_JSON_READ_BYTES:
            checker_summary = {"_error": "summary.json 过大，请在本地打开"}
        else:
            try:
                checker_summary = json.loads(summary_path.read_text(encoding="utf-8"))
            except json.JSONDecodeError as e:
                checker_summary = {"_error": f"JSON 解析失败：{e}"}

    fix_summary: Any = None
    if fix_summary_path.is_file():
        if fix_summary_path.stat().st_size > MAX_JSON_READ_BYTES:
            fix_summary = {"_error": "fix_summary.json 过大"}
        else:
            try:
                fix_summary = json.loads(fix_summary_path.read_text(encoding="utf-8"))
            except json.JSONDecodeError as e:
                fix_summary = {"_error": f"JSON 解析失败：{e}"}

    quality_report_count = 0
    if co.is_dir():
        quality_report_count = len(list(co.rglob("*.quality_report.json")))

    review_md_count = 0
    if co.is_dir():
        review_md_count = len(list(co.rglob("*.review_items.md")))

    return {
        "checker_out": str(co.relative_to(root)).replace("\\", "/"),
        "fix_out": str(fo.relative_to(root)).replace("\\", "/"),
        "checker_summary_path": str(summary_path.relative_to(root)).replace("\\", "/")
        if summary_path.is_file()
        else None,
        "fix_summary_path": str(fix_summary_path.relative_to(root)).replace("\\", "/")
        if fix_summary_path.is_file()
        else None,
        "checker_summary": checker_summary,
        "fix_summary": fix_summary,
        "quality_report_count": quality_report_count,
        "review_items_md_count": review_md_count,
    }


@app.get("/api/reports/content")
def api_reports_content(relative: str) -> Any:
    p = safe_relative_path(relative)
    if not p.is_file():
        raise HTTPException(status_code=404, detail="文件不存在")
    size = p.stat().st_size
    if size > MAX_JSON_READ_BYTES:
        raise HTTPException(status_code=413, detail="文件过大，请在 IDE 中打开")
    text = p.read_text(encoding="utf-8", errors="replace")
    suffix = p.suffix.lower()
    if suffix == ".json":
        try:
            return json.loads(text)
        except json.JSONDecodeError:
            return PlainTextResponse(text, media_type="text/plain; charset=utf-8")
    return PlainTextResponse(text, media_type="text/plain; charset=utf-8")


@app.get("/api/reports/html")
def api_reports_html(relative: str) -> HTMLResponse:
    """以 text/html 返回文件内容，供 iframe 渲染页面（与 /api/reports/content 区分）。"""
    p = safe_relative_path(relative)
    if not p.is_file():
        raise HTTPException(status_code=404, detail="文件不存在")
    suffix = p.suffix.lower()
    if suffix not in (".html", ".htm"):
        raise HTTPException(status_code=400, detail="不是 HTML 文件")
    size = p.stat().st_size
    if size > MAX_JSON_READ_BYTES:
        raise HTTPException(status_code=413, detail="文件过大，请在 IDE 中打开")
    text = p.read_text(encoding="utf-8", errors="replace")
    return HTMLResponse(content=text)


@app.get("/api/reports/pdf")
def api_reports_pdf(relative: str) -> FileResponse:
    p = safe_relative_path(relative)
    if not p.is_file():
        raise HTTPException(status_code=404, detail="文件不存在")
    if p.suffix.lower() != ".pdf":
        raise HTTPException(status_code=400, detail="不是 PDF 文件")
    size = p.stat().st_size
    if size > MAX_PREVIEW_BINARY_BYTES:
        raise HTTPException(status_code=413, detail="文件过大，请在本地打开")
    return FileResponse(
        path=str(p.resolve()),
        media_type="application/pdf",
        filename=p.name,
        content_disposition_type="inline",
    )


@app.get("/api/reports/download")
def api_reports_download(relative: str) -> FileResponse:
    p = safe_relative_path(relative)
    if not p.is_file():
        raise HTTPException(status_code=404, detail="文件不存在")
    size = p.stat().st_size
    if size > MAX_DOWNLOAD_BYTES:
        raise HTTPException(status_code=413, detail="文件过大")
    suffix = p.suffix.lower()
    media = "application/octet-stream"
    if suffix == ".pdf":
        media = "application/pdf"
    elif suffix == ".docx":
        media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif suffix == ".doc":
        media = "application/msword"
    return FileResponse(
        path=str(p.resolve()),
        media_type=media,
        filename=p.name,
        content_disposition_type="attachment",
    )


@app.get("/api/reports/word-preview")
def api_reports_word_preview(relative: str) -> Dict[str, Any]:
    p = safe_relative_path(relative)
    if not p.is_file():
        raise HTTPException(status_code=404, detail="文件不存在")
    suf = p.suffix.lower()
    if suf not in (".doc", ".docx"):
        raise HTTPException(status_code=400, detail="不是 Word 文档")
    size = p.stat().st_size
    if size > MAX_PREVIEW_BINARY_BYTES:
        raise HTTPException(status_code=413, detail="文件过大，请在本地打开")

    root = _root_resolved()
    rel_path = str(p.relative_to(root)).replace("\\", "/")
    download_url = "/api/reports/download?relative=" + quote(rel_path, safe="")

    if suf == ".docx":
        try:
            text = _extract_docx_text(p)
            return {"ok": True, "format": "docx", "text": text}
        except Exception as e:
            return {
                "ok": False,
                "format": "docx",
                "message": f"无法读取该 .docx：{e}",
                "download_url": download_url,
            }

    text = _extract_legacy_doc_text(p)
    if text and text.strip():
        return {"ok": True, "format": "doc", "text": text.strip()}

    return {
        "ok": False,
        "format": "doc",
        "message": "无法在服务端提取该 .doc 的正文。可将 LibreOffice 安装并加入 PATH 后重试，或用 Word 另存为 .docx 再预览。",
        "download_url": download_url,
    }


@app.get("/api/browse")
def api_browse(relative: str = "extract_jsons", q: str = "") -> Dict[str, Any]:
    base = safe_relative_path(relative)
    if not base.exists():
        raise HTTPException(status_code=404, detail="路径不存在")
    if not base.is_dir():
        raise HTTPException(status_code=400, detail="不是目录")
    root = _root_resolved()
    parent_rel: Optional[str] = None
    if base.resolve() != root:
        try:
            parent_rel = str(base.parent.relative_to(root)).replace("\\", "/")
        except ValueError:
            parent_rel = None

    qn = (q or "").strip().lower()

    items: List[Dict[str, Any]] = []
    try:
        for child in sorted(base.iterdir(), key=lambda p: (not p.is_dir(), p.name.lower())):
            try:
                rel = child.relative_to(root)
            except ValueError:
                continue
            entry: Dict[str, Any] = {
                "name": child.name,
                "type": "dir" if child.is_dir() else "file",
                "relative": str(rel).replace("\\", "/"),
            }
            if child.is_file():
                try:
                    entry["size_bytes"] = child.stat().st_size
                except OSError:
                    entry["size_bytes"] = None
            if qn and qn not in entry["name"].lower():
                continue
            items.append(entry)
    except OSError as e:
        raise HTTPException(status_code=500, detail=str(e)) from e
    base_rel = str(base.relative_to(root)).replace("\\", "/")
    return {"base": base_rel, "parent": parent_rel, "items": items, "filter": q or ""}


@app.get("/api/workspace/download-zip")
def api_workspace_download_zip(relative: str) -> FileResponse:
    """将 extract_jsons 下的单个子文件夹打成 zip 下载（不含整个 extract_jsons 根）。"""
    root = _root_resolved()
    base = safe_relative_path(relative)
    _assert_extract_jsons_zip_folder(base, root)

    total = 0
    for fp in base.rglob("*"):
        if fp.is_file():
            try:
                total += fp.stat().st_size
            except OSError:
                continue
            if total > MAX_FOLDER_ZIP_BYTES:
                raise HTTPException(
                    status_code=413,
                    detail=f"文件夹总体积超过上限（约 {MAX_FOLDER_ZIP_BYTES // (1024 * 1024)} MB）",
                )

    td = Path(tempfile.mkdtemp())
    archive_base = td / "pack"
    out_zip = str(archive_base) + ".zip"
    try:
        shutil.make_archive(str(archive_base), "zip", root_dir=str(base.resolve()))
    except OSError as e:
        shutil.rmtree(td, ignore_errors=True)
        raise HTTPException(status_code=500, detail=f"打包失败：{e}") from e
    if not os.path.isfile(out_zip):
        shutil.rmtree(td, ignore_errors=True)
        raise HTTPException(status_code=500, detail="打包未生成文件")

    def _cleanup_pack() -> None:
        shutil.rmtree(td, ignore_errors=True)

    return FileResponse(
        path=out_zip,
        media_type="application/zip",
        filename=f"{base.name}.zip",
        background=BackgroundTask(_cleanup_pack),
    )


@app.post("/api/workspace/delete")
def api_workspace_delete(body: WorkspaceDeleteBody) -> Dict[str, Any]:
    """删除仓库内的文件或文件夹（不可恢复）。禁止删除仓库根与 .git。"""
    root = _root_resolved()
    rel_in = (body.relative or "").strip()
    p = safe_relative_path(rel_in)
    _assert_workspace_delete_allowed(p, root)
    if not p.exists():
        raise HTTPException(status_code=404, detail="路径不存在")
    rel_out = str(p.relative_to(root)).replace("\\", "/")
    try:
        if p.is_dir():
            shutil.rmtree(p)
        else:
            p.unlink()
    except OSError as e:
        raise HTTPException(status_code=500, detail=f"删除失败：{e}") from e
    return {"ok": True, "deleted": rel_out}


@app.get("/api/browse/tree")
def api_browse_tree(relative: str = "extract_jsons", max_depth: int = 3) -> Dict[str, Any]:
    """浅层树（节点数有上限），便于一眼浏览 extract_jsons 结构。"""
    base = safe_relative_path(relative)
    if not base.exists() or not base.is_dir():
        raise HTTPException(status_code=404, detail="路径不存在或不是目录")
    root = _root_resolved()
    max_depth = max(1, min(max_depth, 5))
    counter = [0]

    def walk(d: Path, depth: int) -> List[Dict[str, Any]]:
        out: List[Dict[str, Any]] = []
        if depth >= max_depth or counter[0] >= MAX_TREE_NODES:
            return out
        try:
            kids = sorted(d.iterdir(), key=lambda p: (not p.is_dir(), p.name.lower()))
        except OSError:
            return out
        for child in kids:
            if counter[0] >= MAX_TREE_NODES:
                break
            counter[0] += 1
            try:
                rel = child.relative_to(root)
            except ValueError:
                continue
            node: Dict[str, Any] = {
                "name": child.name,
                "type": "dir" if child.is_dir() else "file",
                "relative": str(rel).replace("\\", "/"),
            }
            if child.is_file():
                try:
                    node["size_bytes"] = child.stat().st_size
                except OSError:
                    node["size_bytes"] = None
            if child.is_dir() and depth + 1 < max_depth:
                node["children"] = walk(child, depth + 1)
            out.append(node)
        return out

    nodes = walk(base, 0)
    base_rel = str(base.relative_to(root)).replace("\\", "/")
    parent_rel: Optional[str] = None
    if base.resolve() != root:
        try:
            parent_rel = str(base.parent.relative_to(root)).replace("\\", "/")
        except ValueError:
            parent_rel = None
    return {
        "base": base_rel,
        "parent": parent_rel,
        "max_depth": max_depth,
        "nodes": nodes,
        "truncated": counter[0] >= MAX_TREE_NODES,
        "node_count": counter[0],
    }
