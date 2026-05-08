"""
将 Markdown 正文渲染为版式化的 Word（python-docx）。
还原 mock_templates 内常见语法：标题 #、无序/有序列表、表格、引用 >、分隔线 ---、行内 **加粗**。

解析逻辑与同包内 `markdown_structure`、PDF 渲染共用，保证 doc / pdf 结构一致。
"""

from __future__ import annotations

import io
import re
from typing import Any, Dict, List

from .markdown_structure import parse_markdown_to_blocks

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    Document = None  # type: ignore


def _set_document_theme(doc: "Document") -> None:
    """正文默认中文雅黑、英文 Calibri；段落适度行距。"""
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(4)


def _add_inline_runs(paragraph, text: str) -> None:
    """解析 **加粗**，其余原样输出。"""
    if not text:
        return
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def _paragraph_bullet(doc: "Document", text: str) -> None:
    try:
        p = doc.add_paragraph(style="List Bullet")
    except (KeyError, ValueError):
        p = doc.add_paragraph()
        p.add_run("• ")
    _add_inline_runs(p, text)


def _paragraph_number(doc: "Document", text: str) -> None:
    try:
        p = doc.add_paragraph(style="List Number")
    except (KeyError, ValueError):
        p = doc.add_paragraph()
    _add_inline_runs(p, text)


def _add_horizontal_rule(doc: "Document") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("―" * 36)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(160, 160, 160)


def _add_markdown_table(doc: "Document", rows: List[List[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    nrows = len(rows)
    table = doc.add_table(rows=nrows, cols=ncols)
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            txt = row[ci] if ci < len(row) else ""
            cell = table.rows[ri].cells[ci]
            p = cell.paragraphs[0]
            p.text = ""
            _add_inline_runs(p, txt)


def _apply_blocks_to_doc(doc: "Document", blocks: List[Dict[str, Any]]) -> None:
    for b in blocks:
        t = b["type"]
        if t == "hr":
            _add_horizontal_rule(doc)
        elif t == "heading":
            doc.add_heading(b["text"], level=min(int(b["level"]), 9))
        elif t == "paragraph":
            p = doc.add_paragraph()
            _add_inline_runs(p, str(b["text"]))
        elif t == "bullet":
            _paragraph_bullet(doc, str(b["text"]))
        elif t == "numbered":
            _paragraph_number(doc, str(b["text"]))
        elif t == "table":
            _add_markdown_table(doc, b["rows"])
        elif t == "quote":
            qp = doc.add_paragraph()
            try:
                qp.style = "Quote"
            except (KeyError, ValueError):
                qp.paragraph_format.left_indent = Inches(0.35)
                qp.paragraph_format.space_before = Pt(4)
                qp.paragraph_format.space_after = Pt(4)
            _add_inline_runs(qp, str(b["text"]))


def render_markdown_to_docx_bytes(document_title: str, md_text: str) -> bytes:
    """
    document_title：封面标题（如【模拟资料】…）。
    md_text：完整 Markdown 正文（含模板与页脚）。
    """
    if Document is None:
        raise ImportError("需要 python-docx：pip install python-docx")

    doc = Document()
    _set_document_theme(doc)

    t = doc.add_heading(document_title, level=0)
    t.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    blocks = parse_markdown_to_blocks(md_text)
    _apply_blocks_to_doc(doc, blocks)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def has_python_docx() -> bool:
    return Document is not None
